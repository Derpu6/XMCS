import streamlit as st
import dashscope
from dashscope import Generation
from langchain_core.messages import AIMessage, HumanMessage, SystemMessage
from langchain_core.outputs import ChatGeneration, ChatResult
from langchain_core.language_models import BaseChatModel
import time
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class QwenChat(BaseChatModel):
    api_key: str

    def _llm_type(self) -> str:
        return "qwen-plus"

    def _generate(self, messages, stop=None, run_manager=None, **kwargs):
        # 将 LangChain 消息格式转换为 DashScope 格式
        dashscope_messages = []
        for msg in messages:
            if isinstance(msg, SystemMessage):
                dashscope_messages.append({"role": "system", "content": msg.content})
            elif isinstance(msg, HumanMessage):
                dashscope_messages.append({"role": "user", "content": msg.content})
            elif isinstance(msg, AIMessage):
                dashscope_messages.append({"role": "assistant", "content": msg.content})
            else:
                dashscope_messages.append({"role": "user", "content": str(msg)})

        try:
            # 构建并发送API请求
            response = Generation.call(
                model="qwen-plus",
                messages=dashscope_messages,
                temperature=0.3,
                api_key=self.api_key
            )

            # 增加健壮性检查
            if response is None:
                raise Exception("API 调用返回 None。请检查网络连接和API密钥。")

            if response.status_code != 200:
                error_msg = f"API Error ({response.status_code}): {getattr(response, 'message', 'Unknown error')}"
                raise Exception(error_msg)

            output = getattr(response, 'output', None)
            if output is None:
                raise Exception("API 响应缺少 'output' 字段。")

            # 从API响应中提取生成的文本内容
            content = None
            if hasattr(output, 'text') and output.text:
                content = output.text
            elif isinstance(output, dict) and 'choices' in output:
                choices = output.get('choices')
                if choices and isinstance(choices, list) and len(choices) > 0:
                    first_choice = choices[0]
                    if isinstance(first_choice, dict) and 'message' in first_choice:
                        message_dict = first_choice['message']
                        if isinstance(message_dict, dict) and 'content' in message_dict:
                            content = message_dict['content']
            if content is None:
                raise Exception("无法从 API 响应中提取生成的文本内容。")

            # 构造 LangChain 兼容的 ChatResult 对象
            message = AIMessage(content=content)
            generation = ChatGeneration(message=message)
            usage = getattr(response, 'usage', {})
            llm_output = {"token_usage": usage, "model_name": self._llm_type()}
            return ChatResult(generations=[generation], llm_output=llm_output)

        except Exception as e:
            raise Exception(f"调用 Qwen 模型时发生错误: {e}")


def create_word_document(content, theme):
    """将文本内容转换为格式化的Word文档[1,6](@ref)"""
    doc = Document()

    # 设置文档标题[6](@ref)
    title = doc.add_paragraph()
    title_run = title.add_run(f"{theme}项目设计方案")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 添加空行
    doc.add_paragraph()

    # 处理内容并添加到文档
    current_paragraph = None

    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue

        # 检测标题格式
        if line.startswith('##### '):
            # 添加标题[6](@ref)
            heading_text = line.replace('##### ', '').strip()
            heading = doc.add_paragraph()
            heading_run = heading.add_run(heading_text)
            heading_run.font.size = Pt(14)
            heading_run.font.bold = True
            current_paragraph = None

        elif line.startswith('        ') or any(
                line.startswith(char) for char in ['-', '*', '•', '1.', '2.', '3.', '4.']):
            # 列表项或缩进内容[6](@ref)
            if current_paragraph is None:
                current_paragraph = doc.add_paragraph()
            else:
                current_paragraph = doc.add_paragraph()

            list_text = line.strip()
            # 移除列表标记前的空格
            if list_text.startswith(('- ', '* ', '• ')):
                list_text = list_text[2:]
            elif any(list_text.startswith(f"{i}.") for i in range(1, 10)):
                list_text = list_text[list_text.find('.') + 1:].strip()

            current_paragraph.add_run("    " + list_text)

        else:
            # 普通段落[6](@ref)
            current_paragraph = doc.add_paragraph()
            current_paragraph.add_run(line)

    return doc


def download_word_file(doc, theme):
    """创建Word文档下载按钮[1](@ref)"""
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)

    st.download_button(
        label="📥 下载Word文档",
        data=bio,
        file_name=f"{theme}项目设计方案.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        help="点击下载完整项目设计方案Word文档",
        key=f"download_{theme}"  # 确保每次生成都有唯一的key
    )


# Streamlit应用
st.title("智能项目功能生成器")

# 在侧边栏设置API密钥
api_key = st.sidebar.text_input("输入API密钥", type="password", help="从阿里云DashScope平台获取")

# 项目设置区域 - 添加多模块选择
st.sidebar.header("项目设置")

# 多模块选择 - 显示屏作为固定选项，但用户可以选择是否包含
modules = st.sidebar.multiselect(
    "选择项目模块",
    ["电机", "显示屏", "串口通信", "外部中断", "定时器"],
    default=["显示屏"],
)

# 确保至少选择了一个模块
if not modules:
    st.sidebar.error("请至少选择一个项目模块")

theme = st.sidebar.text_input("项目主题", "智能流水线")
function = st.sidebar.text_area("项目功能", "自动启停", height=100)


def _get_general_info(modules, theme, function):
    modules_desc = "、".join(modules)

    key_requirements = """
    **特别注意（基于嵌入式课堂要求）：**
    1. 按键功能明确：
        - KEY0: 确定/启动按键
        - KEY1: 调速/调方向/暂停按键（外部中断触发）
        - KEY_UP: 计件传感器/模式切换
    2. 显示要求：
        - 上电初始界面显示系统名称和欢迎标语
        - 主界面显示当前工作模式和运行参数
        - 模式切换时更新主界面上显示的工作模式和运行参数
    3. 功能要求：
        - 支持暂停/恢复功能（按键中断处理）
    4. 计时器使用：
        - TIM3实现1.5秒定时和按键长按检测
        - TIM2实现计数功能
    """

    if "电机" in modules:
        motor_note = "电机模块使用简单转子马达，只需两个信号线控制电压方向，无需使能控制。电机通过pwm调速。"
    else:
        motor_note = ""

    if "显示屏" in modules:
        display_note = "显示屏要求：\n" \
                       "  - 初始界面：第1行显示系统名称，第2行显示欢迎标语\n" \
                       "  - 主界面：第1行显示'Mode:[模式名]'，第2行显示运行参数"
    else:
        display_note = ""

    if "定时器" in modules:
        timer_note = "定时器配置：\n" \
                     "  - TIM3: 实现1.5秒定时，用于刷新界面和状态检测\n" \
                     "  - TIM2: 实现计数功能，记录关键参数"
    else:
        timer_note = ""

    if "外部中断" in modules:
        interrupt_note = "外部中断：\n" \
                         "  - KEY1按键使用外部中断触发暂停功能\n" \
                         "  - 暂停时系统完全停止，界面不刷新\n" \
                         "  - 恢复时保持原有状态继续运行"
    else:
        interrupt_note = ""

    return {
        "modules_desc": modules_desc,
        "motor_note": motor_note,
        "display_note": display_note,
        "timer_note": timer_note,
        "interrupt_note": interrupt_note
    }


def _get_mode_description(theme):
    if any(kw in theme for kw in ["流水线", "传送带", "生产线", "装配线"]):
        return """
            1. **手动控制模式**
               - 用户直接控制设备运行
               - 短按KEY1：循环调节电机速度（30%/50%/70%/90%四档）
               - 短按KEY_UP：切换设备运行方向
               - 指示灯：蓝色闪烁

            2. **自动调节模式**
               - 智能优化生产效率
               - 系统每5秒自动计算最优工作参数
               - KEY_UP作为传感器使用
               - 指示灯：绿色常亮
            """, """
            3. **工作控制**
               - 手动模式：用户根据需求调整工作参数
               - 自动模式：系统基于传感器数据自动优化工作效率
            """

    elif any(kw in theme for kw in ["洗衣机", "搅拌机", "旋转设备", "脱水机"]):
        return """
            1. **低速模式**
               - 轻柔工作状态
               - 低转速（300-500 RPM）
               - 适合精细处理
               - 指示灯：绿色慢闪

            2. **标准模式**
               - 正常工作状态
               - 标准转速（500-800 RPM）
               - 平衡效率和能耗
               - 指示灯：白色常亮

            3. **高速模式**
               - 高效工作状态
               - 高转速（1000-1200 RPM）
               - 最大生产效率
               - 指示灯：红色闪烁
            """, """
            3. **工作控制**
               - 不同模式对应预设工作参数
               - 系统自动应用最合适的运行策略
            """

    elif any(kw in theme for kw in ["票", "售票", "购票", "验票", "选座位"]):
        return """
            1. **售票模式**
               - 主工作状态
               - 处理票务销售
               - 支持多种票型和支付方式
               - 指示灯：绿色常亮

            2. **查询模式**
               - 信息查询状态
               - 查看历史销售记录
               - 分析票务数据
               - 指示灯：蓝色闪烁

            3. **维护模式**
               - 系统维护状态
               - 管理员配置系统参数
               - 系统升级和维护
               - 指示灯：红色闪烁
            """, """
            3. **工作控制**
               - 售票模式：完成票务处理流程
               - 查询模式：提供信息查询服务
               - 维护模式：系统维护和设置
            """

    elif any(kw in theme for kw in ["电梯", "升降机", "垂直运输"]):
        return """
            1. **标准运行模式**
               - 正常工作状态
               - 响应楼层呼叫
               - 优化运输效率
               - 指示灯：白色常亮

            2. **节能运行模式**
               - 低功耗运行状态
               - 减少空载运行
               - 延长设备寿命
               - 指示灯：绿色慢闪

            3. **高峰运行模式**
               - 高效运输状态
               - 优先响应高流量楼层
               - 最大化运输效率
               - 指示灯：红色闪烁

            4. **维护模式**
               - 系统维护状态
               - 工程师调试和检修
               - 指示灯：黄色闪烁
            """, """
            3. **工作控制**
               - 标准模式：平衡效率和能耗
               - 节能模式：低功耗运行
               - 高峰模式：最大化运输效率
               - 维护模式：系统调试和检修
            """

    elif any(kw in theme for kw in ["家居", "智能家居", "家庭自动化"]):
        return """
            1. **居家模式**
               - 家庭成员在家状态
               - 舒适环境设置
               - 灯光和温度自动调节
               - 指示灯：白色常亮

            2. **离家模式**
               - 家庭成员外出状态
               - 节能安全设置
               - 关闭非必要设备
               - 指示灯：蓝色慢闪

            3. **睡眠模式**
               - 夜间休息状态
               - 安静舒适环境
               - 调暗灯光降低噪音
               - 指示灯：紫色慢闪

            4. **娱乐模式**
               - 家庭娱乐状态
               - 优化影音设备
               - 调整灯光氛围
               - 指示灯：彩色循环
            """, """
            3. **工作控制**
               - 居家模式：舒适生活环境
               - 离家模式：节能安全防护
               - 睡眠模式：安静休息环境
               - 娱乐模式：家庭娱乐优化
            """

    elif any(kw in theme for kw in ["农业", "温室", "大棚", "种植"]):
        return """
            1. **自动灌溉模式**
               - 智能浇水状态
               - 根据土壤湿度调节
               - 指示灯：蓝色极亮

            2. **通风模式**
               - 空气循环状态
               - 调节温湿度
               - 防止病虫害
               - 指示灯：绿色慢闪

            极. **补光模式**
               - 光照增强状态
               - 阴天或夜间补充光照
               - 促进植物生长
               - 指示灯：黄色闪烁

            4. **监控极式**
               - 环境监测状态
               - 实时采集温湿度数据
               - 生成生长报告
               - 指示灯：白色闪烁
            """, """
            3. **工作控制**
               - 灌溉模式：智能水分管理
               - 通风模式：优化空气循环
               - 补光模式：补充光照需求
               - 监控模式：环境数据采集
            """

    elif any(kw in theme for kw in ["停车场", "车库", "车位管理"]):
        return """
            1. **入场模式**
               - 车辆进入管理
               - 车牌识别
               - 车位分配
               - 指示灯：绿色常亮

            2. **出场模式**
               - 车辆离开管理
               - 费用计算
               - 支付处理
              极 指示灯：蓝色闪烁

            3. **寻车模式**
               - 帮助车主找车
               - 车位导航
               - 最短路径规划
               - 指示灯：黄色闪烁

            4. **维护模式**
               - 系统维护状态
               - 设备检修
               - 数据备份
               - 指示灯：红色闪烁
            """, """
            3. **工作控制**
               - 入场模式：车辆进入管理
               - 出场模式：车辆离开处理
               - 寻车模式：车位导航服务
               - 维护模式：系统检修维护
            """

    elif any(kw in theme for kw in ["照明", "灯光", "路灯"]):
        return """
            1. **标准照明模式**
               - 正常工作状态
               - 固定亮度设置
               - 指示灯：白色常亮

            2. **节能模式**
               - 低功耗极行状态
               - 根据环境光调节亮度
               - 减少能耗
               - 指示灯：绿色慢闪

            3. **场景模式**
               - 特殊场景设置
               - 如聚会、阅读、观影等
               - 自定义灯光效果
               - 指示灯：彩色循环

            4. **安全模式**
               - 应急照明状态
               - 断电时自动启用
               - 提供基本照明
               - 指示灯：红色闪烁
            """, """
            3. **工作控制**
               - 标准模式：常规照明需求
               - 节能模式：智能亮度调节
               - 场景模式：特殊场景灯光
               - 安全模式：应急照明保障
            """

    else:
        # 通用默认模式方案
        return """
            1. **手动模式**
               - 标准工作状态
               - 通过按键调整工作参数及模式
               - 适合常规工作场景
               - 指示灯：白色常亮

            2. **自动模式**
               - 通过定时器自动调整工作参数及模式
               - 指示灯：绿色慢闪

            """, """
            3. **工作控制**
               - 使用KEY_UP在预设模式间切换
               - KEY1用于微调工作参数
               - 系统根据场景自动优化配置
            """


def create_prompt(modules, theme, function):
    # 获取通用信息
    gen_info = _get_general_info(modules, theme, function)

    mode_desc, work_control = _get_mode_description(theme)

    # 构建完整提示词
    prompt = f"""
        你是一位嵌入式系统课程设计出题专家，请根据嵌入式课堂项目设计标准格式，编写一个基于{gen_info['modules_desc']}的{theme}极目任务书。

        **项目主题：**
        {theme}

        **主要功能：**
        {function}

        **特殊要求：**
        {gen_info['motor_note']}
        {gen_info['display_note']}
        {gen_info['timer_note']}
        {gen_info['interrupt_note']}

        **请严格按照以下结构生成内容（使用中文，不使用Markdown）：**

        ##### 一、任务题目
        一种智能{theme}控制系统

        ##### 二、控制功能
        [用150-200字描述项目整体功能，突出核心控制逻辑和用户交互]

        ##### 三、按键功能
        | 按键名称        | 功能描述                                     |
        |----------     |----------|
        | KEY0          | 确定/启动按键<br>- 短按：启动/恢复控制<br>- 长按(2秒)：进入模式切换 |
        | KEY1          | 多功能按键<br>- 短按：参数调节/功能选择<br>- 长按(>2秒)：暂停/恢复 |
        | KEY_UP        | 传感器/选择器<br>- 正常工作时：信号输入<br>- 模式切换时：选项选择 |

        ##### 四、工作模式
        {mode_desc}

        ##### 五、具体控制要求
        1. **初始界面显示**
           - 上电后在显示屏第1行显示系统名称"{theme}"
           - 第2行显示欢迎标语
           - 短按KEY0后进入主控界面

        2. **模式切换控制**
           - 使用KEY_UP在切换不同模式
           - 显示屏显示新的工作模式和运行参数

        {work_control}

        4. **暂停控制**
           - 系统运行中按下KEY1立即暂停
           - 界面冻结不再刷新
           - 短按KEY0恢复暂停前状态继续运行

        ##### 六、硬件引脚分配
        | 引脚号 | 功能描述 | 类型 | 关联模块 | 用途说明 |
        |--------|----------|------|----------|----------|
        | ...    | ...      | ...  | ...      | ...      |

        **设计约束：**
        1. 只使用预定义的三个按键（KEY0/KEY1/KEY_UP）和两个LED（LED0/LED1）
        2. 硬件引脚分配不包含显示屏相关引脚（如果使用显示屏）
        3. 不包含电源、工作电压等基本因素
        4. 按键功能明确无冲突
    """
    return prompt


# 生成项目功能
def generate_project(prompt, api_key):
    if not api_key:
        return "错误: 请先输入API密钥"

    try:
        qwen = QwenChat(api_key=api_key)
        messages = [
            SystemMessage(content="你是一名经验丰富的嵌入式系统工程师，擅长设计多模块硬件教学系统。"),
            HumanMessage(content=prompt)
        ]
        response = qwen.invoke(messages)
        return response.content
    except Exception as e:
        return f"生成失败: {str(e)}"


# 主界面
st.write("### 项目功能生成器")
st.markdown("选择项目和模块后，点击按钮生成完整设计方案")

if st.button("生成设计方案", type="primary", help="点击生成完整设计文档", disabled=not modules):
    if not api_key:
        st.error("请先输入API密钥")
    elif not modules:
        st.error("请至少选择一个项目模块")
    else:
        progress_bar = st.progress(0)
        status_text = st.empty()

        for percent in range(10):
            progress_bar.progress(percent * 10)
            status_text.text(f"准备项目参数... {percent * 10}%")
            time.sleep(0.05)

        prompt = create_prompt(modules, theme, function)
        progress_bar.progress(30)
        status_text.text("创建设计方案结构...")

        progress_bar.progress(50)
        status_text.text("正在生成设计方案...")
        project_content = generate_project(prompt, api_key)
        progress_bar.progress(90)

        if project_content.startswith("错误:") or project_content.startswith("生成失败:"):
            status_text.error(project_content)
        else:
            progress_bar.progress(100)
            status_text.success(f"{theme}项目设计方案生成成功！")

            st.subheader(f"{theme}项目设计方案")
            st.markdown(project_content, unsafe_allow_html=True)

            # 添加Word文档下载功能
            try:
                doc = create_word_document(project_content, theme)
                download_word_file(doc, theme)
                st.success("Word文档已准备就绪，请点击上方下载按钮保存")
            except Exception as e:
                st.error(f"创建Word文档时出错: {str(e)}")

        progress_bar.empty()
